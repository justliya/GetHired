from google.adk.agents.llm_agent import Agent
from google.adk.agents import SequentialAgent
from google.adk.tools.load_artifacts_tool import load_artifacts_tool
from google.adk.tools.tool_context import ToolContext

from docx import Document
from . import prompt
import io
SECTION_HEADERS = {
    'summary': ['summary', 'objective', 'about me'],
    'experience': ['experience', 'work history', 'professional experience'],
    'skills': ['skills', 'technologies', 'technical skills'],
    'education': ['education', 'academic background'],
    'certifications': ['certifications', 'licenses'],
    'projects': ['projects', 'portfolio', 'selected projects']
}


def load_resume_job_desc():
    doc = Document('docs/LaKaleigh_Harris_Resume.docx')
    doc_content = [p.text.strip() for p in doc.paragraphs]
    with open('docs/job_post.txt', 'r', encoding='utf-8') as f:
        job_content = f.read()
    return doc_content, job_content
def load_template_doc():
    doc = Document()
    
resume, job_spec = load_resume_job_desc()

base_resume_cleanup =  Agent(
    model="gemini-2.0-flash",
    name="base_resume_cleanup",
    instruction = prompt.BASE_PROMPT.format("Lakaleigh","Software Engineer", "AI", resume),
    output_key="base_resume",
    tools=[load_resume_job_desc]
    )

job_optimization_agent =  Agent(
    model="gemini-2.0-flash",
    name="job_optimization_agent",
    instruction = prompt.JOB_OPTIMIZATION.format(job_spec),
    output_key="job_optimized_resume"
    )

experience_optimization_agent  =  Agent(
    model="gemini-2.0-flash",
    name="experience_optimization_agent",
    instruction = prompt.EXPERIENCE_OPTIMIZATION.format(job_spec),
    output_key='experience_optimizated_resume'
    )

ats_optimization_agent =  Agent(
    model="gemini-2.0-flash",
    name="ats_optimization_agent",
    instruction = prompt.ATS_OPTIMIZATION.format(job_spec),
    output_key='ats_optimized_resume'
    )

humanize_resume_agent =  Agent(
    model="gemini-2.0-flash",
    name="humanize_resume_agent",
    instruction = prompt.HUMANIZE.format(job_spec),
    output_key='draft_resume'
    )

proof_reader_agent =  Agent(
    model="gemini-2.0-flash",
    name="humanize_resume_agent",
    instruction = prompt.PROOF_READ_RESUME,
    output_key='final_resume'
    )

resume_pipeline_agent = SequentialAgent(
    name='ResumeTailorAgent',
    description='Agent design to tailor resumes to a job description',
    sub_agents=[
        base_resume_cleanup,
        job_optimization_agent,
        experience_optimization_agent,
        ats_optimization_agent,
        humanize_resume_agent,
        proof_reader_agent
        ],
)
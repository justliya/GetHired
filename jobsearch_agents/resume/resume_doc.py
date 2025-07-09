import os
import uuid
import tempfile
import logging
import re
import json
from datetime import datetime, timedelta
from pathlib import Path
from docxtpl import DocxTemplate
from .parse import ParsedResume
from typing import Optional
import firebase_admin
from firebase_admin import credentials, storage

try:
    import requests
except ImportError:
    requests = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

logger = logging.getLogger(__name__)

# Global constant for invalid user IDs
INVALID_USER_IDS = {
    # Literal strings
    "user_id", "USER_ID", "UserId", "userId", "user-id",
    
    # Common placeholders
    "ACTUAL_USER_ID", "actual_user_id", "ACTUAL-USER-ID",
    "YOUR_USER_ID", "your_user_id", "YOUR-USER-ID",
    "REPLACE_WITH_USER_ID", "INSERT_USER_ID_HERE",
    
    # Template variables
    "{user_id}", "{{user_id}}", "${user_id}", "%(user_id)s",
    "<user_id>", "[user_id]", "[[user_id]]", ":user_id",
    "[USER_ID]", "<USER_ID>", "{USER_ID}",
    
    # Programming artifacts
    "null", "None", "undefined", "nil", "NULL",
    "false", "true", "0", "-1",
    
    # Empty values
    "", " ", "  ", "\n", "\t",
    
    # Common test values
    "test", "TEST", "test_user", "demo", "DEMO",
    "example", "EXAMPLE", "sample", "SAMPLE",
    
    # Placeholder patterns
    "XXX", "xxx", "TODO", "FIXME", "TBD",
    "placeholder", "PLACEHOLDER", "dummy", "DUMMY",
}


def initialize_firebase():
    """Initialize Firebase Admin SDK if not already initialized"""
    try:
        firebase_admin.get_app()
        logger.debug("Firebase already initialized")
    except ValueError:
        # No app initialized yet
        SERVICE_ACCOUNT_KEY_PATH = os.getenv('SERVICE_ACCOUNT_KEY_PATH')
        FIREBASE_STORAGE_BUCKET = os.getenv('FIREBASE_STORAGE_BUCKET')
        
        if SERVICE_ACCOUNT_KEY_PATH and os.path.exists(SERVICE_ACCOUNT_KEY_PATH):
            cred = credentials.Certificate(SERVICE_ACCOUNT_KEY_PATH)
            firebase_admin.initialize_app(cred, {
                'storageBucket': FIREBASE_STORAGE_BUCKET
            })
            logger.info("Firebase initialized with service account for bucket: %s", FIREBASE_STORAGE_BUCKET)
        else:
            # Initialize with default credentials
            firebase_admin.initialize_app(options={
                'storageBucket': FIREBASE_STORAGE_BUCKET
            })
            logger.info("Firebase initialized with default credentials for bucket: %s", FIREBASE_STORAGE_BUCKET)


def is_valid_user_id(user_id: Optional[str]) -> bool:
    """Check if a user_id is valid (not a placeholder)
    
    Firebase Authentication UIDs are typically:
    - 28 characters long
    - Alphanumeric with possible special characters
    - Example: Gn8mXRcszzOPvGIYomUmHWMxA0E2
    """
    if not user_id or not isinstance(user_id, str):
        return False
    
    normalized = str(user_id).strip()
    
    # Check against invalid list
    if normalized.lower() in {s.lower() for s in INVALID_USER_IDS}:
        return False
    
    # Check length - Firebase UIDs are typically 28 chars but can vary
    if len(normalized) < 10 or len(normalized) > 128:
        return False
    
    # Check for placeholder patterns
    placeholder_patterns = [
        r'^user.?id$',  # Matches user_id, user-id, user.id, etc.
        r'^\{.*\}$',    # Matches anything in curly braces
        r'^<.*>$',      # Matches anything in angle brackets
        r'^\$\{.*\}$',  # Matches ${...}
    ]
    
    for pattern in placeholder_patterns:
        if re.match(pattern, normalized, re.IGNORECASE):
            return False
    
    # Check if it contains only valid characters
    if not re.match(r'^[a-zA-Z0-9_\-]+$', normalized):
        return False
    
    return True


def create_unique_filename(job_position_title: str, user_id: str) -> str:
    """Create a unique, secure, easily-lookupable resume file name"""
    safe_job_title = "".join(c for c in job_position_title if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_job_title = safe_job_title.replace(' ', '_')[:50]
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    unique_id = str(uuid.uuid4())[:8]
    
    filename = f"resume_{safe_job_title}_{timestamp}_{unique_id}.docx"
    logger.debug("Created filename: %s for user: %s", filename, user_id)
    
    return filename


def extract_user_id_from_url(resume_url: str) -> str:
    """Extract user_id from Firebase Storage URL
    
    Expected format:
    https://storage.googleapis.com/[project-id].firebasestorage.app/resumes/[user_id]/[filename]
    
    Example:
    https://storage.googleapis.com/gethired-6c623.firebasestorage.app/resumes/Gn8mXRcszzOPvGIYomUmHWMxA0E2/resume_AI_Research_Engineer_20250708_210856_992270fc.docx
    
    Where user_id = Gn8mXRcszzOPvGIYomUmHWMxA0E2
    
    Returns:
        str: Extracted user_id or empty string if not found
    """
    try:
        if not resume_url:
            logger.warning("❌ No resume URL provided for extraction")
            return ""
        
        logger.debug("🔍 Attempting to extract user_id from URL: %s", resume_url)
        
        # Primary pattern for the exact Firebase Storage format
        # Matches: storage.googleapis.com/[project].firebasestorage.app/resumes/[user_id]/[filename]
        primary_pattern = r'storage\.googleapis\.com/[^/]+\.firebasestorage\.app/resumes/([^/]+)/'
        
        match = re.search(primary_pattern, resume_url)
        if match:
            extracted_id = match.group(1)
            logger.info("✅ Extracted user_id: '%s' from Firebase Storage URL", extracted_id)
            
            # Validate the extracted ID
            if is_valid_user_id(extracted_id):
                return extracted_id
            else:
                logger.warning("❌ Extracted value '%s' appears to be invalid", extracted_id)
                return ""
        
        # Fallback patterns for other possible formats
        fallback_patterns = [
            r'/resumes/([^/]+)/',                    # Generic pattern
            r'resumes%2F([^%]+)%2F',                 # URL encoded version
        ]
        
        for pattern in fallback_patterns:
            match = re.search(pattern, resume_url)
            if match:
                extracted_id = match.group(1)
                logger.info("✅ Extracted user_id: '%s' using fallback pattern: %s", extracted_id, pattern)
                
                if is_valid_user_id(extracted_id):
                    return extracted_id
                    
        logger.warning("❌ Could not extract valid user_id from URL: %s", resume_url)
        return ""
        
    except Exception as e:
        logger.error("❌ Error extracting user_id from URL %s: %s", resume_url, e)
        return ""


def upload_to_firebase_storage(file_path: str, filename: str, user_id: str) -> dict:
    """Upload file to Firebase Storage and return download URL
    
    Creates URL in format:
    https://storage.googleapis.com/[project].firebasestorage.app/resumes/[user_id]/[filename]
    """
    try:
        logger.info("🚀 Starting Firebase Storage upload process")
        logger.info("  📁 File path: %s", file_path)
        logger.info("  📄 Filename: %s", filename)
        logger.info("  👤 User ID: %s", user_id)
        
        # Validate user_id format
        if not is_valid_user_id(user_id):
            logger.error("❌ Invalid user_id format: '%s'", user_id)
            raise ValueError(f"Invalid user_id: {user_id}")
        
        # Validate file exists and has content
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Source file does not exist: {file_path}")
        
        file_size = os.path.getsize(file_path)
        logger.info("  📊 File size: %d bytes", file_size)
        
        if file_size == 0:
            raise ValueError(f"Source file is empty: {file_path}")
        
        # Initialize Firebase if not already done
        initialize_firebase()
        
        # Get bucket
        bucket = storage.bucket()
        logger.info("  🪣 Using Firebase Storage bucket: %s", bucket.name)
        
        # Create storage path: resumes/[user_id]/[filename]
        storage_path = f"resumes/{user_id}/{filename}"
        logger.info("  🗂️  Storage path: %s", storage_path)
        
        # Create blob
        blob = bucket.blob(storage_path)
        
        # Set metadata
        blob.metadata = {
            'uploadedBy': user_id,
            'uploadTime': datetime.now().isoformat(),
            'fileType': 'tailored_resume',
            'originalSize': str(file_size),
            'contentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        # Upload file
        logger.info("  ⬆️  Starting file upload...")
        blob.upload_from_filename(file_path, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        logger.info("  ✅ File upload completed successfully")
        
        # Make the blob publicly accessible
        logger.info("  🌐 Setting blob to public access...")
        blob.make_public()
        
        # Construct the correct Firebase Storage URL format
        # Format: https://storage.googleapis.com/[bucket_name]/resumes/[user_id]/[filename]
        public_url = f"https://storage.googleapis.com/{bucket.name}/resumes/{user_id}/{filename}"
        
        # Generate a signed URL for authenticated access (7 days validity as mentioned in template)
        signed_url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(days=7),
            method="GET"
        )
        
        logger.info("✅ Successfully uploaded resume to Firebase Storage")
        logger.info("  📥 Public URL: %s", public_url)
        logger.info("  🔐 Signed URL: %s", signed_url[:100] + "...")  # Log first 100 chars
        
        return {
            "success": True,
            "download_url": public_url,
            "public_url": public_url,
            "signed_url": signed_url,
            "storage_path": storage_path,
            "bucket": bucket.name,
            "firebase_url": public_url  # Same as public_url for this format
        }
        
    except Exception as e:
        logger.error("❌ Error during Firebase Storage upload:")
        logger.error("  🚨 Error type: %s", type(e).__name__)
        logger.error("  🚨 Error details: %s", e)
        logger.exception("  📋 Full stack trace:")
        
        return {
            "success": False,
            "error": f"Upload error: {str(e)}",
            "download_url": "",
            "public_url": "",
            "firebase_url": "",
            "signed_url": ""
        }


def create_formatted_resume(text: str, job_position_title: str = "Position", user_id: str = 'user_id', resume_url: Optional[str] = None) -> dict:
    """Create a formatted resume document and upload to Firebase Storage
    
    Args:
        text: The actual
CORRECT USAGE EXAMPLES:

1. With Firebase authenticated user:
   result = create_formatted_resume(
       text=resume_text,
       job_position_title="AI Research Engineer",
       user_id=" Gn8mXRcszzOPvGIYomUmHWMxA0E2"  # Real Firebase UID
   )

2. With URL fallback:
   result = create_formatted_resume(
       text=resume_text,
       job_position_title="Software Engineer",
       resume_url="https://storage.googleapis.com/gethired-6c623.firebasestorage.app/resumes/ Gn8mXRcszzOPvGIYomUmHWMxA0E2/original.pdf"
   )

3. Testing functions:
   # Test URL extraction
   test_firebase_url_extraction()
   
   # Test user ID validation
   validate_user_id_handling()
   
   # Test Firebase connectivity
   test_firebase_connectivity()
"""
def create_formatted_resume(text: str, job_position_title: str = "Position", user_id: str = 'user_id', resume_url: Optional[str] = None) -> dict:
    """Create a formatted resume document and upload to Firebase Storage
    
    Args:
        text: The actual resume text content
        job_position_title: The job position title (e.g., "Software Engineer")
        user_id: The ACTUAL user ID value (e.g., "Gn8mXRcszzOPvGIYomUmHWMxA0E2")
                 NOT the literal string "user_id" or placeholder text!
        resume_url: Optional resume URL for user_id extraction fallback
    
    Returns:
        dict: Result with URLs and status
    """
    temp_file_path = None
    
    try:
        # === CRITICAL USER_ID VALIDATION ===
        logger.info("🔍 VALIDATING USER_ID...")
        logger.info("  📥 Received user_id parameter: '%s'", user_id)
        logger.info("  📥 Received resume_url parameter: '%s'", resume_url)
        
        # Normalize the user_id for checking
        normalized_user_id = str(user_id).strip() if user_id else ""
        
        # Check if user_id is valid
        if not is_valid_user_id(normalized_user_id):
            logger.error("❌ INVALID USER_ID DETECTED: '%s'", user_id)
            
            if normalized_user_id.lower() in {s.lower() for s in INVALID_USER_IDS}:
                logger.error("   This appears to be a placeholder or literal string, not an actual user ID!")
                logger.error("   Expected format like: 'Gn8mXRcszzOPvGIYomUmHWMxA0E2'")
            
            # Try to extract from URL
            if resume_url:
                logger.info("🔧 Attempting to extract user_id from resume URL...")
                extracted_user_id = extract_user_id_from_url(resume_url)
                if extracted_user_id and is_valid_user_id(extracted_user_id):
                    user_id = extracted_user_id
                    logger.info("✅ Successfully extracted valid user_id from URL: '%s'", user_id)
                else:
                    logger.error("❌ Could not extract valid user_id from URL")
                    user_id = f"anonymous_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
                    logger.warning("⚠️ Using generated anonymous ID: '%s'", user_id)
            else:
                user_id = f"anonymous_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
                logger.warning("⚠️ No valid user_id or resume_url provided, using generated ID: '%s'", user_id)
        else:
            user_id = normalized_user_id
            logger.info("✅ Valid user_id provided: '%s'", user_id)
            
        # Log final decision
        logger.info("📌 FINAL USER_ID TO BE USED: '%s'", user_id)
        logger.info("=" * 60)
        
        # Validate resume text
        if not text or text.strip() == "":
            error_msg = "Resume text is empty or None"
            logger.error("❌ %s", error_msg)
            return {
                "resume_text": text,
                "download_url": "",
                "public_url": "",
                "signed_url": "",
                "filename": "",
                "status": "error",
                "message": error_msg
            }
        
        logger.info("🚀 Starting resume creation for user '%s', position: '%s'", user_id, job_position_title)
        logger.info("📄 Input resume text length: %d characters", len(text))
        logger.debug("📋 First 500 chars of resume text: %s", text[:500])
        
        # Parse the resume text
        logger.info("🔍 Parsing resume text...")
        try:
            parsed_resume = ParsedResume(text)
            candidate_data = parsed_resume.serialize()
            logger.info("✅ Resume parsed successfully")
        except Exception as parse_error:
            logger.error("❌ Failed to parse resume: %s", parse_error)
            return {
                "resume_text": text,
                "download_url": "",
                "public_url": "",
                "signed_url": "",
                "filename": "",
                "status": "error",
                "message": f"Failed to parse resume: {str(parse_error)}"
            }
        
        # Log the candidate data for debugging
        candidate = candidate_data.get("candidate", {})
        logger.info("📋 Parsed candidate info:")
        logger.info("  - Name: %s", candidate.get("name"))
        logger.info("  - Email: %s", candidate.get("email"))
        logger.info("  - Phone: %s", candidate.get("phone"))
        logger.info("  - Education entries: %d", len(candidate.get("education", [])))
        logger.info("  - Experience entries: %d", len(candidate.get("experience_section", [])))
        
        # Log skills structure for debugging
        skills_data = candidate.get("skills", {})
        if isinstance(skills_data, dict):
            logger.info("  - Technical skills: %d", len(skills_data.get("technical", [])))
            logger.info("  - Soft skills: %d", len(skills_data.get("soft_skills", [])))
        else:
            logger.info("  - Skills count: %d", len(skills_data) if isinstance(skills_data, list) else 0)
        
        # Find template file
        possible_paths = [
            Path("/app/template/templateResumeDocV2.docx"),
            Path(__file__).parent.parent / "template" / "templateResumeDocV2.docx",
            Path(__file__).parent.parent.parent / "template" / "templateResumeDocV2.docx",
        ]
        
        template_path = None
        for path in possible_paths:
            logger.debug("🔍 Checking template path: %s", path)
            if path.exists():
                template_path = path
                logger.info("✅ Found template at: %s", template_path)
                break
        
        if not template_path:
            raise FileNotFoundError(f"Template file not found. Tried paths: {[str(p) for p in possible_paths]}")
            
        # Create the DocxTemplate and render
        logger.info("🔧 Loading DocxTemplate...")
        doc = DocxTemplate(str(template_path))
        
        # The template expects everything under 'candidate'
        # Don't modify the structure - just pass it through
        template_context = {
            "candidate": candidate_data["candidate"]
        }
        
        # Handle skills structure for the template
        # The template expects either technical_skills/soft_skills or skills_flat
        if "skills" in candidate:
            skills = candidate["skills"]
            if isinstance(skills, dict):
                # If skills is a dict with technical/soft_skills keys
                if "technical" in skills:
                    candidate["technical_skills"] = skills["technical"]
                if "soft_skills" in skills:
                    candidate["soft_skills"] = skills["soft_skills"]
            elif isinstance(skills, list):
                # If skills is just a flat list
                candidate["skills_flat"] = skills
        
        # Debug: Check if the data structure matches template expectations
        logger.info("=== TEMPLATE CONTEXT VALIDATION ===")
        expected_fields = [
            "name", "email", "phone", "professional_summary", 
            "education", "experience_section", "projects", 
            "certifications", "volunteer_experience"
        ]
        
        for field in expected_fields:
            if field in candidate:
                field_value = candidate[field]
                if isinstance(field_value, list):
                    logger.info("✓ Field '%s' present: list with %d items", field, len(field_value))
                elif isinstance(field_value, str):
                    logger.info("✓ Field '%s' present: string (%d chars)", field, len(field_value))
                else:
                    logger.info("✓ Field '%s' present: %s", field, type(field_value).__name__)
            else:
                logger.warning("✗ Field '%s' missing", field)
        
        # Log the final context being sent
        logger.info("📋 Final template context being sent:")
        logger.info(json.dumps(template_context, indent=2, default=str))
        
        logger.info("🎨 Rendering template with context...")
        try:
            doc.render(template_context)
            logger.info("✅ Template rendered successfully")
        except Exception as render_error:
            logger.error("❌ Template rendering failed: %s", render_error)
            logger.error("Context that failed: %s", json.dumps(template_context, indent=2, default=str))
            raise
        
        # Create filename
        filename = create_unique_filename(job_position_title, user_id)
        logger.info("📁 Generated filename: %s", filename)
        
        # Create temporary file and save document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file_path = temp_file.name
            logger.info("📂 Temporary file path: %s", temp_file_path)
            
            # Save the document
            doc.save(temp_file_path)
            logger.info("✅ Document saved successfully")
        
        # Verify the file was created
        if os.path.exists(temp_file_path):
            file_size = os.path.getsize(temp_file_path)
            logger.info("📊 Generated document size: %d bytes", file_size)
        else:
            raise FileNotFoundError(f"Temporary document file was not created: {temp_file_path}")
        
        # Upload to Firebase Storage
        logger.info("☁️  Uploading to Firebase Storage...")
        upload_result = upload_to_firebase_storage(temp_file_path, filename, user_id)
        
        if not upload_result.get("success", False):
            logger.warning("⚠️  Failed to upload to storage")
            return {
                "resume_text": text,
                "download_url": "",
                "public_url": "",
                "signed_url": "",
                "filename": filename,
                "status": "error",
                "message": f"Resume created but upload failed: {upload_result.get('error', 'Unknown error')}"
            }
        
        logger.info("🎉 Resume created successfully for %s (user: %s)", job_position_title, user_id)
        
        return {
            "resume_text": text,
            "download_url": upload_result.get("download_url", ""),
            "public_url": upload_result.get("public_url", ""),
            "signed_url": upload_result.get("signed_url", ""),
            "filename": filename,
            "storage_path": upload_result.get("storage_path", ""),
            "bucket": upload_result.get("bucket", ""),
            "status": "success",
            "message": f"Resume successfully created and uploaded for {job_position_title}"
        }
        
    except Exception as e:
        error_message = f"Unexpected error creating formatted resume: {str(e)}"
        logger.error("❌ %s", error_message)
        logger.exception("Full error details:")
        
        return {
            "resume_text": text,
            "document_url": "",
            "download_url": "",
            "public_url": "",
            "signed_url": "",
            "filename": "",
            "status": "error",
            "message": error_message
        }
        
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.debug("🧹 Cleaned up temporary file: %s", temp_file_path)
            except OSError as e:
                logger.warning("⚠️  Failed to delete temporary file %s: %s", temp_file_path, e)


def download_and_extract_resume_text(resume_url: str) -> str:
    """
    Download a resume file from Firebase Storage, extract text content, and clean up
    
    Args:
        storage_url: Firebase Storage URL in format:
        https://storage.googleapis.com/[project].firebasestorage.app/resumes/[user_id]/[filename]
        
    Returns:
        str: Extracted text content from the resume file
        
    Raises:
        ValueError: If file download fails or text extraction fails
    """
    temp_file_path = None
    
    try:
        if requests is None:
            raise ImportError("requests package not installed. Install with: pip install requests")
            
        if DocxDocument is None:
            logger.warning("python-docx package not installed. DOCX file processing may not work.")
            
        if PyPDF2 is None:
            logger.warning("PyPDF2 package not installed. PDF file processing may not work.")
        
        # Download the file
        logger.info("Downloading resume from: %s", resume_url)
        response = requests.get(resume_url, timeout=30)
        response.raise_for_status()
        
        # Determine file type from URL or content type
        file_extension = None
        if resume_url.lower().endswith('.docx'):
            file_extension = '.docx'
        elif resume_url.lower().endswith('.pdf'):
            file_extension = '.pdf'
        elif resume_url.lower().endswith('.txt'):
            file_extension = '.txt'
        else:
            # Try to determine from content type
            content_type = response.headers.get('content-type', '').lower()
            if 'word' in content_type or 'officedocument' in content_type:
                file_extension = '.docx'
            elif 'pdf' in content_type:
                file_extension = '.pdf'
            elif 'text' in content_type:
                file_extension = '.txt'
            else:
                logger.warning("Unknown file type, defaulting to .docx")
                file_extension = '.docx'
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file_path = temp_file.name
            temp_file.write(response.content)
        
        logger.debug("Downloaded file to temporary location: %s", temp_file_path)
        
        # Extract text based on file type
        if file_extension == '.docx':
            if DocxDocument is None:
                raise ValueError("error processing please try again")
                
            doc = DocxDocument(temp_file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            extracted_text = '\n'.join(text_content)
            
        elif file_extension == '.pdf':
            if PyPDF2 is None:
                raise ValueError("PyPDF2 package not installed. Cannot process PDF files.")
                
            with open(temp_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_content = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text.strip())
                
                extracted_text = '\n'.join(text_content)
                
        elif file_extension == '.txt':
            with open(temp_file_path, 'r', encoding='utf-8') as txt_file:
                extracted_text = txt_file.read()
                
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not extracted_text.strip():
            raise ValueError("No text content could be extracted from the file")
        
        logger.info("Successfully extracted %d characters from resume", len(extracted_text))
        return extracted_text.strip()
        
    except requests.RequestException as e:
        error_msg = f"Failed to download file from {resume_url}: {str(e)}"
        logger.error(error_msg)
        raise ValueError(error_msg)
        
    except Exception as e:
        error_msg = f"Unexpected error extracting text from resume: {str(e)}"
        logger.error(error_msg)
        raise ValueError(error_msg)
        
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.debug("Cleaned up temporary download file: %s", temp_file_path)
            except OSError as e:
                logger.warning("Failed to delete temporary download file %s: %s", temp_file_path, e)


def test_firebase_connectivity(bucket_name: Optional[str] = None) -> dict:
    """Test Firebase Storage connectivity and bucket access"""
    try:
        if not bucket_name:
            bucket_name = os.getenv('FIREBASE_STORAGE_BUCKET')
        
        logger.info("🧪 Testing Firebase Storage connectivity...")
        logger.info("  🪣 Target bucket: %s", bucket_name)
        
        # Initialize Firebase
        initialize_firebase()
        
        # Get bucket
        bucket = storage.bucket()
        logger.info("  🪣 Got bucket: %s", bucket.name)
        
        # Test bucket access by creating a test file
        test_blob_name = f"test_connectivity_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        test_blob = bucket.blob(test_blob_name)
        test_content = f"Test connectivity at {datetime.now().isoformat()}"
        
        logger.info("  ✍️  Testing write access with test file: %s", test_blob_name)
        test_blob.upload_from_string(test_content, content_type='text/plain')
        
        # Verify the test file was uploaded
        logger.info("  🔍 Verifying test file upload...")
        if test_blob.exists():
            logger.info("  ✅ Write test successful - file exists")
            
            # Test public access
            try:
                test_blob.make_public()
                public_url = test_blob.public_url
                logger.info("  🌐 Public URL: %s", public_url)
                public_access = True
            except Exception as e:
                logger.warning("  ⚠️  Could not make blob public: %s", e)
                public_access = False
            
            # Clean up test file
            test_blob.delete()
            logger.info("  🧹 Test file cleaned up")
            
            return {
                "success": True,
                "message": "Firebase Storage connectivity test passed",
                "bucket_name": bucket.name,
                "public_access": public_access
            }
        else:
            return {
                "success": False,
                "error": "Test file was not created successfully",
                "bucket_name": bucket.name
            }
            
    except Exception as e:
        logger.error("❌ Firebase Storage connectivity test failed: %s", e)
        logger.exception("Full error details:")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }


def test_firebase_url_extraction():
    """Test extraction from the specific Firebase Storage URL format"""
    test_cases = [
        (
            "https://storage.googleapis.com/gethired-6c623.firebasestorage.app/resumes/Gn8mXRcszzOPvGIYomUmHWMxA0E2/resume_AI_Research_Engineer_20250708_210856_992270fc.docx",
            "n8mXRcszzOPvGIYomUmHWMxA0E2",
            "Standard Firebase Storage URL"
        ),
        (
            "https://storage.googleapis.com/myproject-12345.firebasestorage.app/resumes/5f4dcc3b5aa765d61d8327deb882/resume_Software_Engineer_20250115_143022_a1b2c3d4.docx",
            "5f4dcc3b5aa765d61d8327deb882",
            "Another valid Firebase URL"
        ),
        (
            "https://storage.googleapis.com/test-app.firebasestorage.app/resumes/user_id/resume.docx",
            "",
            "Invalid - contains placeholder 'user_id'"
        ),
        (
            "https://storage.googleapis.com/app.firebasestorage.app/resumes/john_doe_123/document.pdf",
            "john_doe_123",
            "Username format user ID"
        ),
    ]
    
    logger.info("🧪 Testing Firebase Storage URL extraction...")
    for url, expected_id, description in test_cases:
        extracted = extract_user_id_from_url(url)
        status = "✅" if extracted == expected_id else "❌"
        logger.info("%s %s", status, description)
        logger.info("   URL: %s", url[:80] + "..." if len(url) > 80 else url)
        logger.info("   Expected: '%s', Got: '%s'", expected_id, extracted)


def validate_user_id_handling():
    """Test function to validate user_id handling"""
    test_cases = [
        # (user_id_input, expected_valid, description)
        ("Gn8mXRcszzOPvGIYomUmHWMxA0E2", True, "Valid Firebase user ID"),
        ("5f4dcc3b5aa765d61d8327deb882", True, "Valid Firebase user ID (28 chars)"),
        ("john_doe_123", True, "Valid username format"),
        ("abc123def456", True, "Valid alphanumeric ID"),
        ("user_id", False, "Literal string 'user_id'"),
        ("ACTUAL_USER_ID", False, "Placeholder text"),
        ("{user_id}", False, "Template variable"),
        ("[USER_ID]", False, "Bracket placeholder"),
        ("", False, "Empty string"),
        (None, False, "None value"),
    ]
    
    logger.info("🧪 Running user_id validation tests...")
    for user_id, expected_valid, description in test_cases:
        is_valid = is_valid_user_id(user_id)
        status = "✅" if is_valid == expected_valid else "❌"
        logger.info("%s %s: '%s' -> Valid: %s (Expected: %s)", 
                   status, description, user_id, is_valid, expected_valid)


# Example usage documentation
"""
# Example 1: Create resume with valid user ID
result = create_formatted_resume(
    text="John Doe\nSoftware Engineer\n...",
    job_position_title="Senior Software Engineer",
    user_id="Gn8mXRcszzOPvGIYomUmHWMxA0E2"  

# Example 2: Create resume with URL fallback for user ID extraction
result = create_formatted_resume(
    text="Jane Smith\nData Scientist\n...",
    job_position_title="Data Scientist",
    user_id="user_id",  # Invalid placeholder
    resume_url="https://storage.googleapis.com/gethired.firebasestorage.app/resumes/abc123def456/resume.docx"
)

# Example 3: Test Firebase connectivity
test_result = test_firebase_connectivity()
"""